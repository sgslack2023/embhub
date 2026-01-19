import os
import tempfile
import re
import requests
import logging
from datetime import datetime
from io import BytesIO
from django.http import HttpResponse
from rest_framework.decorators import action
import openpyxl
from rest_framework.views import APIView
from rest_framework.viewsets import ModelViewSet
from rest_framework.response import Response
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
from back_sinan.custom_methods import isAuthenticatedCustom
from .google_drive_service import GoogleDriveService
from .track123_service import import_tracking_to_track123, get_tracking_status
from users.models import GoogleDriveSettings, UserActivities
from .models import Product, Account, PackingSlip, File
from .serializers import ProductSerializer, AccountSerializer, PackingSlipSerializer, FileSerializer, ExpenseSerializer

logger = logging.getLogger(__name__)

# PDF generation imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
from PIL import Image as PILImage
import fitz  # PyMuPDF for PDF processing

class EMBHubDriveAccountsView(APIView):
    """Get available Google Drive accounts"""
    permission_classes = (isAuthenticatedCustom,)

    def get(self, request):
        try:
            accounts = GoogleDriveService.get_available_drive_accounts()
            return Response({
                'success': True,
                'accounts': accounts
            })
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubFolderTreeView(APIView):
    """Get folder tree structure"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            folder_id = request.data.get('folder_id')  # Optional, if None will get root
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            tree = service.get_folder_tree(folder_id)
            
            return Response({
                'success': True,
                'tree': tree
            })
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubCreateFolderView(APIView):
    """Create a new folder"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            name = request.data.get('name')
            parent_id = request.data.get('parent_id')
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not name or not name.strip():
                return Response({
                    'error': 'Folder name is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Validate folder name
            invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
            if any(char in name for char in invalid_chars):
                return Response({
                    'error': 'Folder name contains invalid characters'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            folder = service.create_folder(name.strip(), parent_id)
            
            return Response({
                'success': True,
                'folder': folder,
                'message': f'Folder "{name}" created successfully'
            })
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubDeleteFolderView(APIView):
    """Delete a folder"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            folder_id = request.data.get('folder_id')
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not folder_id:
                return Response({
                    'error': 'Folder ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            success = service.delete_folder(folder_id)
            
            if success:
                return Response({
                    'success': True,
                    'message': 'Folder deleted successfully'
                })
            else:
                return Response({
                    'error': 'Failed to delete folder'
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubUploadFileView(APIView):
    """Upload a file to a folder"""
    permission_classes = (isAuthenticatedCustom,)
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        print(f"=== EMBHubUploadFileView called ===")
        print(f"Request method: {request.method}")
        print(f"Request FILES: {list(request.FILES.keys())}")
        
        try:
            email = request.data.get('google_drive_email')
            folder_id = request.data.get('folder_id')
            uploaded_file = request.FILES.get('file')
            
            print(f"Email: {email}")
            print(f"Folder ID: {folder_id}")
            print(f"Uploaded file: {uploaded_file.name if uploaded_file else None}")
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not folder_id:
                return Response({
                    'error': 'Folder ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not uploaded_file:
                return Response({
                    'error': 'File is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Check file size (max 100MB)
            max_size = 100 * 1024 * 1024  # 100MB
            if uploaded_file.size > max_size:
                return Response({
                    'error': 'File size must be less than 100MB'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            try:
                service = GoogleDriveService(email)
                uploaded_file_info = service.upload_file(
                    temp_file_path, 
                    uploaded_file.name, 
                    folder_id,
                    uploaded_file.content_type
                )
                
                # Check if this is a Word file uploaded to a "Packing Slips" folder
                packing_slips_processing = self.process_packing_slips_if_applicable(
                    service, folder_id, uploaded_file, temp_file_path
                )
                
                # Check if this is a PDF file uploaded to a "Shipping Labels" folder
                shipping_labels_processing = self.process_shipping_labels_if_applicable(
                    service, folder_id, uploaded_file, temp_file_path, uploaded_file_info
                )
                
                # Check if this is a DST/DGT file uploaded to a "DST" folder
                dst_dgt_processing = self.process_dst_dgt_files_if_applicable(
                    service, folder_id, uploaded_file, uploaded_file_info
                )
                
                response_data = {
                    'success': True,
                    'file': uploaded_file_info,
                    'message': f'File "{uploaded_file.name}" uploaded successfully'
                }
                
                # Add packing slips processing results if applicable
                if packing_slips_processing:
                    response_data.update(packing_slips_processing)
                
                # Add shipping labels processing results if applicable
                if shipping_labels_processing:
                    response_data.update(shipping_labels_processing)
                
                # Add DST/DGT processing results if applicable
                if dst_dgt_processing:
                    response_data.update(dst_dgt_processing)
                
                return Response(response_data)
            finally:
                # Clean up temporary file (ignore Windows file locking errors)
                if os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except OSError:
                        # Windows file locking - OS will clean up temp files eventually
                        pass
                
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

    def process_packing_slips_if_applicable(self, service, folder_id, uploaded_file, temp_file_path):
        """Process packing slips if Word file is uploaded to Packing Slips folder"""
        try:
            print(f"=== PROCESSING PACKING SLIPS CHECK ===")
            print(f"File: {uploaded_file.name}")
            print(f"Folder ID: {folder_id}")
            print(f"Temp file path: {temp_file_path}")
            
            # Check if file is a Word document
            if not uploaded_file.name.lower().endswith(('.doc', '.docx')):
                print(f"File {uploaded_file.name} is not a Word document, skipping packing slip processing")
                return None
            
            print(f"✓ File {uploaded_file.name} is a Word document, checking folder...")
            
            # Get folder name to check if it's a "Packing Slips" folder
            folder_name = service.get_folder_name(folder_id)
            print(f"Folder name retrieved: '{folder_name}'")
            
            if folder_name.lower() != 'packing slips':
                print(f"✗ Folder '{folder_name}' is not 'packing slips', skipping processing")
                return None
            
            print(f"✓ File uploaded to Packing Slips folder, starting processing...")
            
            # Get full folder path for tracking
            try:
                folder_path_list = service.get_folder_path(folder_id)
                # Convert list of folder objects to readable path string
                folder_path = " / ".join([folder['name'] for folder in folder_path_list])
                print(f"Full folder path: '{folder_path}'")
            except Exception as e:
                print(f"Error getting folder path: {e}")
                folder_path = folder_name  # Fallback to just folder name
            
            # Process the Word document
            print(f"Parsing Word document from: {temp_file_path}")
            parsed_labels = self.parse_word_document_from_path(temp_file_path)
            print(f"Parsed {len(parsed_labels) if parsed_labels else 0} labels from document")
            
            if not parsed_labels:
                return {
                    'packing_slips_processing': {
                        'processed': False,
                        'message': 'No valid packing slips found in the document'
                    }
                }
            
            # Save to database
            print(f"Saving {len(parsed_labels)} labels to database...")
            created_labels = []
            errors = []
            
            for i, label_data in enumerate(parsed_labels):
                print(f"Processing label {i+1}: {label_data}")
                # Add folder path to the label data
                label_data['folder_path'] = folder_path
                print(f"Label data with folder_path: {label_data}")
                serializer = PackingSlipSerializer(data=label_data)
                if serializer.is_valid():
                    packing_slip = serializer.save()
                    created_labels.append({
                        'id': packing_slip.id,
                        'order_id': packing_slip.order_id,
                        'product_code': packing_slip.product.code,
                        'product_name': packing_slip.product.name
                    })
                    print(f"Successfully saved label {i+1}: Order {packing_slip.order_id}")
                else:
                    print(f"Validation errors for label {i+1}: {serializer.errors}")
                    print(f"Failed data: {label_data}")
                    errors.append({
                        'data': label_data,
                        'errors': serializer.errors
                    })
            
            print(f"Database save complete: {len(created_labels)} created, {len(errors)} errors")
            
            return {
                'packing_slips_processing': {
                    'processed': True,
                    'created_labels': len(created_labels),
                    'errors': len(errors),
                    'total_parsed': len(parsed_labels),
                    'labels': created_labels[:10],  # Show first 10 labels
                    'parse_errors': errors[:3] if errors else []  # Show first 3 errors
                }
            }
            
        except Exception as e:
            print(f"=== ERROR IN PACKING SLIPS PROCESSING ===")
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'packing_slips_processing': {
                    'processed': False,
                    'error': f'Error processing packing slips: {str(e)}'
                }
            }

    def process_shipping_labels_if_applicable(self, service, folder_id, uploaded_file, temp_file_path, uploaded_file_info):
        """Process shipping labels if PDF file is uploaded to Shipping Labels folder"""
        try:
            print(f"=== PROCESSING SHIPPING LABELS CHECK ===")
            print(f"File: {uploaded_file.name}")
            print(f"Folder ID: {folder_id}")
            print(f"Temp file path: {temp_file_path}")
            
            # Check if file is a PDF document
            if not uploaded_file.name.lower().endswith('.pdf'):
                print(f"File {uploaded_file.name} is not a PDF document, skipping shipping label processing")
                return None
            
            print(f"✓ File {uploaded_file.name} is a PDF document, checking folder...")
            
            # Get folder name to check if it's a "Shipping Labels" folder
            folder_name = service.get_folder_name(folder_id)
            print(f"Folder name retrieved: '{folder_name}'")
            
            if folder_name.lower() not in ['shipping labels', 'shipping label']:
                print(f"✗ Folder '{folder_name}' is not 'shipping labels' or 'shipping label', skipping processing")
                print(f"Available folder names that would work: 'Shipping Labels' or 'Shipping Label'")
                return None
            
            print(f"✓ File uploaded to Shipping Labels folder, starting processing...")
            
            # Get full folder path for tracking
            try:
                folder_path_list = service.get_folder_path(folder_id)
                # Convert list of folder objects to readable path string
                folder_path = " / ".join([folder['name'] for folder in folder_path_list])
                print(f"Full folder path: '{folder_path}'")
            except Exception as e:
                print(f"Error getting folder path: {e}")
                folder_path = folder_name  # Fallback to just folder name
            
            # Get all packing slips for matching based on folder path
            # Extract parent folder path (excluding "Shipping Labels" folder)
            parent_folder_path = self.extract_parent_folder_path(folder_path)
            print(f"Parent folder path for matching: '{parent_folder_path}'")
            
            # Get packing slips that share the same parent folder path
            packing_slips = list(PackingSlip.objects.filter(
                folder_path__icontains=parent_folder_path
            ).values('id', 'ship_to', 'order_id', 'folder_path'))
            print(f"Retrieved {len(packing_slips)} packing slips from matching folder path")
            
            if len(packing_slips) == 0:
                print(f"⚠️ No packing slips found with folder path containing: '{parent_folder_path}'")
                print("Available packing slips with their folder paths:")
                all_packing_slips = PackingSlip.objects.all().values('id', 'order_id', 'folder_path')[:10]
                for ps in all_packing_slips:
                    print(f"  - ID {ps['id']}: {ps['order_id']} -> '{ps['folder_path']}'")
                return {
                    'shipping_labels_processing': {
                        'processed': False,
                        'message': f'No packing slips found in matching folder path: {parent_folder_path}'
                    }
                }
            
            # Process the PDF document
            print(f"Processing PDF document from: {temp_file_path}")
            try:
                from .pdf_utils import PDFProcessor
                processor = PDFProcessor()
                print(f"✓ PDF processor initialized successfully")
            except ImportError as e:
                print(f"✗ Error importing PDFProcessor: {e}")
                return {
                    'shipping_labels_processing': {
                        'processed': False,
                        'error': f'PDF processing libraries not available: {str(e)}'
                    }
                }
            
            # Get the Google Drive file link from the uploaded file info
            google_drive_file_link = uploaded_file_info.get('webViewLink', '')
            print(f"Google Drive file link: {google_drive_file_link}")
            
            processed_labels = processor.process_shipping_labels_pdf(temp_file_path, packing_slips, google_drive_file_link)
            print(f"Processed {len(processed_labels) if processed_labels else 0} shipping labels from PDF")
            
            if not processed_labels:
                return {
                    'shipping_labels_processing': {
                        'processed': False,
                        'message': 'No valid shipping labels found in the PDF'
                    }
                }
            
            # Save shipping label files to database
            print(f"Saving {len(processed_labels)} shipping labels to database...")
            created_files = []
            unmatched_labels = []
            
            for i, label_data in enumerate(processed_labels):
                print(f"Processing shipping label {i+1}: {label_data}")
                
                if label_data['matched']:
                    # Create File record for matched shipping label with Google Drive link and page number
                    file_record = File.objects.create(
                        packing_slip_id=label_data['packing_slip_id'],
                        file_type='shipping_label',
                        file_path=label_data['google_drive_file_link'],
                        page_number=label_data['page_number']
                    )
                    created_files.append({
                        'id': file_record.id,
                        'page_number': label_data['page_number'],
                        'file_path': label_data['google_drive_file_link'],
                        'packing_slip_id': label_data['packing_slip_id'],
                        'confidence_score': label_data['confidence_score']
                    })
                    print(f"Successfully saved shipping label {i+1} for packing slip {label_data['packing_slip_id']}")
                else:
                    unmatched_labels.append({
                        'page_number': label_data['page_number'],
                        'shipping_address': label_data['shipping_address'],
                        'file_path': label_data['google_drive_file_link']
                    })
                    print(f"Shipping label {i+1} could not be matched to any packing slip")
            
            print(f"Database save complete: {len(created_files)} matched, {len(unmatched_labels)} unmatched")
            
            return {
                'shipping_labels_processing': {
                    'processed': True,
                    'matched_labels': len(created_files),
                    'unmatched_labels': len(unmatched_labels),
                    'total_processed': len(processed_labels),
                    'created_files': created_files[:10],  # Show first 10 files
                    'unmatched_labels': unmatched_labels[:3] if unmatched_labels else []  # Show first 3 unmatched
                }
            }
            
        except Exception as e:
            print(f"=== ERROR IN SHIPPING LABELS PROCESSING ===")
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'shipping_labels_processing': {
                    'processed': False,
                    'error': f'Error processing shipping labels: {str(e)}'
                }
            }

    def extract_parent_folder_path(self, full_folder_path):
        """Extract parent folder path by removing the last folder (Packing Slips or Shipping Labels)"""
        try:
            # Split the path by " / " separator
            path_parts = full_folder_path.split(" / ")
            
            # Remove the last part (should be "Packing Slips" or "Shipping Labels")
            if len(path_parts) > 1:
                parent_path = " / ".join(path_parts[:-1])
                print(f"Extracted parent path: '{parent_path}' from '{full_folder_path}'")
                return parent_path
            else:
                # If there's only one part, return it as is
                print(f"Single folder path, returning as is: '{full_folder_path}'")
                return full_folder_path
                
        except Exception as e:
            print(f"Error extracting parent folder path: {e}")
            return full_folder_path

    def process_dst_dgt_files_if_applicable(self, service, folder_id, uploaded_file, uploaded_file_info):
        """Process DST/DGT files if uploaded to DST folder and match with packing slips based on order ID"""
        try:
            print(f"=== PROCESSING DST/DGT FILES CHECK ===")
            print(f"File: {uploaded_file.name}")
            print(f"Folder ID: {folder_id}")
            
            # Check if file has DST or DGT extension
            file_extension = uploaded_file.name.lower().split('.')[-1]
            if file_extension not in ['dst', 'dgt']:
                print(f"File {uploaded_file.name} is not a DST or DGT file, skipping processing")
                return None
            
            print(f"✓ File {uploaded_file.name} is a {file_extension.upper()} file, checking folder...")
            
            # Get folder name to check if it's a "DST" folder
            folder_name = service.get_folder_name(folder_id)
            print(f"Folder name retrieved: '{folder_name}'")
            
            if folder_name.lower() != 'dst':
                print(f"✗ Folder '{folder_name}' is not 'DST', skipping processing")
                return None
            
            print(f"✓ File uploaded to DST folder, starting processing...")
            
            # Extract order ID from filename
            # New format: "Customer Name (order-number).dgt" -> extract "order-number" from parentheses
            # Old format: "113-7500760-1326650.dgt" -> use entire filename as order ID
            filename_without_ext = uploaded_file.name.rsplit('.', 1)[0]
            
            # Try to extract order ID from parentheses first (new format)
            import re
            parentheses_match = re.search(r'\(([^)]+)\)', filename_without_ext)
            if parentheses_match:
                order_id = parentheses_match.group(1).strip()
                print(f"Extracted order ID from parentheses: '{order_id}'")
            else:
                # Fallback to old format (entire filename is order ID)
                order_id = filename_without_ext
                print(f"Using entire filename as order ID: '{order_id}'")
            
            # Find matching packing slip by order ID
            try:
                packing_slip = PackingSlip.objects.get(order_id=order_id)
                print(f"✓ Found matching packing slip with ID: {packing_slip.id}")
                
                # Create File record for DST/DGT file
                file_record = File.objects.create(
                    packing_slip=packing_slip,
                    file_type=file_extension,
                    file_path=uploaded_file_info.get('webViewLink', '')
                )
                
                created_file = {
                    'id': file_record.id,
                    'file_type': file_extension,
                    'file_path': file_record.file_path,
                    'packing_slip_id': packing_slip.id,
                    'order_id': order_id
                }
                
                print(f"Successfully saved {file_extension.upper()} file for packing slip {packing_slip.id}")
                
                return {
                    'dst_dgt_processing': {
                        'processed': True,
                        'file_type': file_extension,
                        'matched': True,
                        'order_id': order_id,
                        'packing_slip_id': packing_slip.id,
                        'created_file': created_file
                    }
                }
                
            except PackingSlip.DoesNotExist:
                print(f"✗ No packing slip found with order ID: '{order_id}'")
                return {
                    'dst_dgt_processing': {
                        'processed': True,
                        'file_type': file_extension,
                        'matched': False,
                        'order_id': order_id,
                        'message': f'No packing slip found with order ID: {order_id}'
                    }
                }
            
        except Exception as e:
            print(f"=== ERROR IN DST/DGT FILES PROCESSING ===")
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'dst_dgt_processing': {
                    'processed': False,
                    'error': f'Error processing DST/DGT file: {str(e)}'
                }
            }

    def parse_word_document_from_path(self, file_path):
        """Parse Word document from file path and extract packing slip information"""
        try:
            from docx import Document
        except ImportError:
            raise Exception("python-docx library is required. Please install it using: pip install python-docx")
        
        # Read the document
        print(f"Opening Word document: {file_path}")
        doc = Document(file_path)
        
        # Extract all text from the document
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Also extract text from tables if any
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
        
        if table_text:
            full_text += "\n--- TABLE CONTENT ---\n" + table_text
            print(f"Found table content: {table_text[:200]}")

        print(f"Extracted text length: {len(full_text)} characters")
        print(f"First 1000 characters: {full_text[:1000]}")
        print(f"Raw text with escaped characters: {repr(full_text[:500])}")

        # Extract shipping addresses from table if available
        shipping_addresses = []
        if table_text:
            # Parse addresses from table - typically name, street, city/state/zip pattern
            address_lines = [line.strip() for line in table_text.split('\n') if line.strip()]
            current_address = []
            
            for line in address_lines:
                current_address.append(line)
                # If line looks like city/state/zip (contains state and zip), it's the end of an address
                if re.search(r'[A-Z]{2}\s+\d{5}', line):
                    if len(current_address) >= 2:  # At least name and one address line
                        shipping_addresses.append('\n'.join(current_address))
                    current_address = []
            
            print(f"Extracted {len(shipping_addresses)} shipping addresses from table:")
            for i, addr in enumerate(shipping_addresses):
                print(f"Address {i+1}: {repr(addr)}")

        # Split by packing slips (looking for "Ship to" pattern)
        label_sections = re.split(r'(?=Ship to)', full_text, flags=re.IGNORECASE)
        print(f"Found {len(label_sections)} sections after splitting by 'Ship to'")
        
        parsed_labels = []
        address_index = 0
        
        for i, section in enumerate(label_sections):
            if not section.strip() or 'ship to' not in section.lower():
                print(f"Section {i} skipped (empty or no 'ship to')")
                continue
            
            print(f"Processing section {i}: {section[:200]}...")
            label_data = self.extract_label_data_from_text(section)
            
            if label_data:
                # Assign shipping address from table if available
                if address_index < len(shipping_addresses):
                    label_data['ship_to'] = shipping_addresses[address_index]
                    print(f"Assigned address {address_index + 1} to section {i}: {repr(label_data['ship_to'])}")
                    address_index += 1
                else:
                    print(f"No more addresses available for section {i}")
                
                print(f"Successfully extracted label data from section {i}")
                parsed_labels.append(label_data)
            else:
                print(f"No label data extracted from section {i}")
        
        return parsed_labels

    def extract_label_data_from_text(self, text):
        """Extract individual shipping label data from text section"""
        try:
            print(f"Extracting label data from text section (length: {len(text)})")
            
            # Initialize data structure
            label_data = {
                'ship_to': '',
                'order_id': '',
                'asin': '',
                'sku': '',  # This will be used to look up the product
                'product': None,  # Will be set after product lookup
                'customizations': '',
                'quantity': 1
            }
            
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            print(f"Split into {len(lines)} non-empty lines")
            
            # Extract shipping address (everything between "Ship to" and "Order ID")
            print(f"Text section being processed: {repr(text[:200])}")
            
            # Try multiple regex patterns
            patterns = [
                r'Ship to\s*\n(.*?)(?=Order ID:)',  # Original pattern
                r'Ship to\s*(.*?)(?=Order ID:)',    # Without mandatory newline
                r'Ship to\s*\n+(.*?)(?=Order ID:)', # With one or more newlines
                r'Ship to\s*[\n\r]+(.*?)(?=Order ID:)', # With various line endings
            ]
            
            ship_to_found = False
            for i, pattern in enumerate(patterns):
                print(f"Trying pattern {i+1}: {pattern}")
                ship_to_match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
                
                if ship_to_match:
                    ship_to_content = ship_to_match.group(1).strip()
                    print(f"Pattern {i+1} matched! Raw content: {repr(ship_to_content)}")
                    # Split into lines and remove empty ones
                    ship_to_lines = [line.strip() for line in ship_to_content.split('\n') if line.strip()]
                    label_data['ship_to'] = '\n'.join(ship_to_lines)
                    print(f"Extracted ship_to using pattern {i+1}: '{label_data['ship_to']}'")
                    ship_to_found = True
                    break
                else:
                    print(f"Pattern {i+1} failed")
            
            if not ship_to_found:
                print("All regex patterns failed, trying line-by-line parsing")
                # Fallback to original line-by-line parsing
                ship_to_lines = []
                in_ship_to = False
                
                print(f"Lines to process: {lines}")
                for i, line in enumerate(lines):
                    print(f"Processing line {i}: '{line}'")
                    if 'ship to' in line.lower():
                        print(f"Found 'ship to' at line {i}")
                        in_ship_to = True
                        continue
                    elif 'order id' in line.lower():
                        print(f"Found 'order id' at line {i}, stopping")
                        break
                    elif in_ship_to and line:
                        print(f"Adding ship_to line: '{line}'")
                        ship_to_lines.append(line)
                
                label_data['ship_to'] = '\n'.join(ship_to_lines)
                print(f"Extracted ship_to using line parsing: '{label_data['ship_to']}'")
            
            # Extract Order ID
            order_match = re.search(r'Order ID:\s*#?\s*([^\s\n]+)', text, re.IGNORECASE)
            if order_match:
                label_data['order_id'] = order_match.group(1).strip()
            
            # Extract ASIN
            asin_match = re.search(r'ASIN:\s*([^\s\n]+)', text, re.IGNORECASE)
            if asin_match:
                label_data['asin'] = asin_match.group(1).strip()
            
            # Extract product code from SKU field and lookup product
            code_match = re.search(r'SKU:\s*([^\s\n]+)', text, re.IGNORECASE)
            if code_match:
                code = code_match.group(1).strip()
                try:
                    product = Product.objects.get(code=code)
                    label_data['product'] = product.id  # Set product ID
                    print(f"Found product with code {code}: {product.name} (ID: {product.id})")
                except Product.DoesNotExist:
                    print(f"No product found with code: {code}")
                    return None
            
            # Extract Quantity
            qty_match = re.search(r'QTY:\s*(\d+)', text, re.IGNORECASE)
            if qty_match:
                label_data['quantity'] = int(qty_match.group(1))
            
            # Extract Customizations
            customization_patterns = [
                r'Customizations:(.*?)(?=QTY:|$)',
                r'Left Chest Customization:(.*?)(?=QTY:|$)',
                r'Surface 1:(.*?)(?=QTY:|$)'
            ]
            
            customizations = []
            for pattern in customization_patterns:
                matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
                for match in matches:
                    customizations.append(match.strip())
            
            if customizations:
                label_data['customizations'] = '\n'.join(customizations)
            
            # Validate required fields
            print(f"Validation - Order ID: '{label_data['order_id']}', Product: '{label_data.get('product')}'")
            if not label_data['order_id'] or not label_data.get('product'):
                print(f"Validation failed: missing Order ID or Product")
                return None
                
            print(f"Label validation passed, returning: {label_data}")
            return label_data
            
        except Exception as e:
            print(f"Error extracting label data: {str(e)}")
            return None


class EMBHubDeleteFileView(APIView):
    """Delete a file"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            file_id = request.data.get('file_id')
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not file_id:
                return Response({
                    'error': 'File ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            success = service.delete_file(file_id)
            
            if success:
                return Response({
                    'success': True,
                    'message': 'File deleted successfully'
                })
            else:
                return Response({
                    'error': 'Failed to delete file'
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubListFolderContentsView(APIView):
    """List contents of a specific folder"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            folder_id = request.data.get('folder_id')  # Optional, if None will get root
            
            print(f"EMBHubListFolderContentsView: email={email}, folder_id={folder_id}")
            
            if not email:
                print("Error: Google Drive email is required")
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            print(f"Initializing GoogleDriveService for email: {email}")
            service = GoogleDriveService(email)
            print("GoogleDriveService initialized successfully")
            
            if folder_id:
                print(f"Loading contents for folder_id: {folder_id}")
                folders = service.list_folders(folder_id)
                files = service.list_files(folder_id)
                path = service.get_folder_path(folder_id)
                print(f"Folder path: {path}")
            else:
                print("Loading root folder contents")
                # Get root folder
                root_folder = service.get_root_folder()
                print(f"Root folder: {root_folder}")
                folders = service.list_folders(root_folder['id'])
                files = service.list_files(root_folder['id'])
                path = [root_folder]
            
            print(f"Found {len(folders)} folders and {len(files)} files")
            print(f"Folders: {[f['name'] for f in folders]}")
            print(f"Files: {[f['name'] for f in files]}")
            
            return Response({
                'success': True,
                'folders': folders,
                'files': files,
                'path': path
            })
        except Exception as e:
            print(f"Error in EMBHubListFolderContentsView: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)

class EMBHubSearchView(APIView):
    """Search for files"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            query = request.data.get('query')
            folder_id = request.data.get('folder_id')  # Optional
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not query or not query.strip():
                return Response({
                    'error': 'Search query is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            results = service.search_files(query.strip(), folder_id)
            
            return Response({
                'success': True,
                'results': results,
                'query': query.strip()
            })
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)


def add_user_activity(user, action):
    """Helper function to log user activities"""
    UserActivities.objects.create(
        user_id=user.id,
        email=user.email,
        fullname=user.fullname,
        action=action
    )


class ProductViewSet(ModelViewSet):
    """
    A viewset that provides default `create()`, `retrieve()`, `update()`,
    `partial_update()`, `destroy()` and `list()` actions for Product model.
    """
    queryset = Product.objects.all()
    serializer_class = ProductSerializer
    permission_classes = (isAuthenticatedCustom,)

    @action(detail=False, methods=['get'], url_path='template')
    def download_template(self, request):
        """Download Excel template for bulk product upload"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Product Template"

        # Define headers
        headers = ['name', 'code', 'sku_description', 'sku_uom', 'sku_buy_cost', 'sku_price', 'color']
        ws.append(headers)

        # Add sample data
        ws.append(['Sample Product', 'SP001', 'Sample Description', 'Each', 10.00, 20.00, 'Red'])

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=product_template.xlsx'
        wb.save(response)
        return response

    @action(detail=False, methods=['post'], url_path='bulk-create')
    def bulk_create_products(self, request):
        """Bulk create products from JSON data"""
        serializer = self.get_serializer(data=request.data, many=True)
        serializer.is_valid(raise_exception=True)
        products = serializer.save()
        
        # Log activity
        add_user_activity(request.user, f"Bulk uploaded {len(products)} products")
        
        return Response({'success': True, 'count': len(products)}, status=status.HTTP_201_CREATED)


    def perform_create(self, serializer):
        """Override to add user activity logging"""
        product = serializer.save()
        add_user_activity(self.request.user, f"added new product: {product.name}")

    def perform_update(self, serializer):
        """Override to add user activity logging"""
        product = serializer.save()
        add_user_activity(self.request.user, f"updated product: {product.name}")

    def perform_destroy(self, instance):
        """Override to add user activity logging"""
        product_name = instance.name
        instance.delete()
        add_user_activity(self.request.user, f"deleted product: {product_name}")


class AccountViewSet(ModelViewSet):
    """
    A viewset that provides default `create()`, `retrieve()`, `update()`,
    `partial_update()`, `destroy()` and `list()` actions for Account model.
    """
    queryset = Account.objects.all()
    serializer_class = AccountSerializer
    permission_classes = (isAuthenticatedCustom,)

    def perform_create(self, serializer):
        """Override to add user activity logging"""
        account = serializer.save()
        add_user_activity(self.request.user, f"added new account: {account.account_name}")

    def perform_update(self, serializer):
        """Override to add user activity logging"""
        account = serializer.save()
        add_user_activity(self.request.user, f"updated account: {account.account_name}")

    def perform_destroy(self, instance):
        """Override to add user activity logging"""
        account_name = instance.account_name
        instance.delete()
        add_user_activity(self.request.user, f"deleted account: {account_name}")


class ExpenseViewSet(ModelViewSet):
    """CRUD operations for expenses"""
    permission_classes = (isAuthenticatedCustom,)
    serializer_class = ExpenseSerializer
    
    def get_queryset(self):
        from .models import Expense
        return Expense.objects.select_related('account').all()
    
    def perform_create(self, serializer):
        """Override to add user activity logging"""
        expense = serializer.save()
        add_user_activity(
            self.request.user, 
            f"created expense: {expense.account.account_name} - {expense.get_expense_type_display()} - ${expense.amount}"
        )
    
    def perform_update(self, serializer):
        """Override to add user activity logging"""
        expense = serializer.save()
        add_user_activity(
            self.request.user, 
            f"updated expense: {expense.account.account_name} - {expense.get_expense_type_display()} - ${expense.amount}"
        )
    
    def perform_destroy(self, instance):
        """Override to add user activity logging"""
        expense_info = f"{instance.account.account_name} - {instance.get_expense_type_display()} - ${instance.amount}"
        instance.delete()
        add_user_activity(self.request.user, f"deleted expense: {expense_info}")


class EMBHubRunAutomationView(APIView):
    """Run automation to create daily folders with account subfolders"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            email = request.data.get('google_drive_email')
            
            if not email:
                return Response({
                    'error': 'Google Drive email is required'
                }, status=status.HTTP_400_BAD_REQUEST)

            service = GoogleDriveService(email)
            
            # Get current date in MM-DD-YYYY format
            current_date = datetime.now().strftime('%m-%d-%Y')
            
            # Get all active accounts
            accounts = Account.objects.filter(active=True)
            
            if not accounts.exists():
                return Response({
                    'error': 'No active accounts found'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Create main date folder
            date_folder = service.create_folder(current_date, None)
            
            created_folders = []
            errors = []
            
            # Create subfolders for each account
            for account in accounts:
                try:
                    # Create account folder
                    account_folder = service.create_folder(account.account_name, date_folder['id'])
                    
                    # Create the 4 required subfolders
                    subfolder_names = ['DST', 'Shipping Labels', 'Packing Slips', 'Pick Lists']
                    account_subfolders = []
                    
                    for subfolder_name in subfolder_names:
                        subfolder = service.create_folder(subfolder_name, account_folder['id'])
                        account_subfolders.append(subfolder['name'])
                    
                    created_folders.append({
                        'account_name': account.account_name,
                        'account_folder_id': account_folder['id'],
                        'subfolders': account_subfolders
                    })
                    
                except Exception as e:
                    errors.append({
                        'account_name': account.account_name,
                        'error': str(e)
                    })
            
            # Log the automation activity
            add_user_activity(
                request.user, 
                f"ran automation for {current_date}: created folders for {len(created_folders)} accounts"
            )
            
            return Response({
                'success': True,
                'message': f'Automation completed successfully for {current_date}',
                'date_folder': {
                    'name': current_date,
                    'id': date_folder['id']
                },
                'created_folders': created_folders,
                'errors': errors,
                'total_accounts': len(accounts),
                'successful_accounts': len(created_folders),
                'failed_accounts': len(errors)
            })
            
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)


class PackingSlipsUploadView(APIView):
    """Upload and parse packing slips from Word documents"""
    permission_classes = (isAuthenticatedCustom,)
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        try:
            file = request.FILES.get('file')
            if not file:
                return Response({
                    'error': 'No file uploaded'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Check if file is a Word document
            if not file.name.lower().endswith(('.doc', '.docx')):
                return Response({
                    'error': 'Only Word documents (.doc, .docx) are supported'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Parse the document
            parsed_labels = self.parse_word_document(file)
            
            if not parsed_labels:
                return Response({
                    'error': 'No valid packing slips found in the document'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Save to database
            created_labels = []
            errors = []
            
            for label_data in parsed_labels:
                # Add folder path for manual uploads (standalone upload)
                label_data['folder_path'] = 'Manual Upload'
                serializer = PackingSlipSerializer(data=label_data)
                if serializer.is_valid():
                    packing_slip = serializer.save()
                    created_labels.append({
                        'id': packing_slip.id,
                        'order_id': packing_slip.order_id,
                        'product_code': packing_slip.product.code,
                        'product_name': packing_slip.product.name
                    })
                else:
                    errors.append({
                        'data': label_data,
                        'errors': serializer.errors
                    })

            return Response({
                'success': True,
                'created_labels': len(created_labels),
                'errors': len(errors),
                'total_parsed': len(parsed_labels),
                'label_ids': created_labels,
                'parse_errors': errors[:5]  # Show first 5 errors
            })

        except Exception as e:
            return Response({
                'error': f'Failed to process file: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def parse_word_document(self, file):
        """Parse Word document and extract packing slip information"""
        try:
            from docx import Document
        except ImportError:
            raise Exception("python-docx library is required. Please install it using: pip install python-docx")
        
        # Read the document
        doc = Document(file)
        
        # Extract all text from the document
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Split by packing slips (looking for "Ship to" pattern)
        label_sections = re.split(r'(?=Ship to)', full_text, flags=re.IGNORECASE)
        
        parsed_labels = []
        
        for section in label_sections:
            if not section.strip() or 'ship to' not in section.lower():
                continue
                
            label_data = self.extract_label_data(section)
            if label_data:
                parsed_labels.append(label_data)
        
        return parsed_labels

    def extract_label_data(self, text):
        """Extract individual packing slip data from text section"""
        try:
            # Initialize data structure
            label_data = {
                'ship_to': '',
                'order_id': '',
                'asin': '',
                'sku': '',  # This will be used to look up the product
                'customizations': '',
                'quantity': 1
            }
            
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            # Extract shipping address (everything between "Ship to" and "Order ID")
            ship_to_lines = []
            in_ship_to = False
            
            for i, line in enumerate(lines):
                if 'ship to' in line.lower():
                    in_ship_to = True
                    continue
                elif 'order id' in line.lower():
                    break
                elif in_ship_to and line:
                    ship_to_lines.append(line)
            
            label_data['ship_to'] = '\n'.join(ship_to_lines)
            
            # Extract Order ID
            order_match = re.search(r'Order ID:\s*#?\s*([^\s\n]+)', text, re.IGNORECASE)
            if order_match:
                label_data['order_id'] = order_match.group(1).strip()
            
            # Extract ASIN
            asin_match = re.search(r'ASIN:\s*([^\s\n]+)', text, re.IGNORECASE)
            if asin_match:
                label_data['asin'] = asin_match.group(1).strip()
            
            # Extract product code from SKU field and lookup product
            code_match = re.search(r'SKU:\s*([^\s\n]+)', text, re.IGNORECASE)
            if code_match:
                code = code_match.group(1).strip()
                try:
                    product = Product.objects.get(code=code)
                    label_data['product'] = product.id  # Set product ID
                    print(f"Found product with code {code}: {product.name} (ID: {product.id})")
                except Product.DoesNotExist:
                    print(f"No product found with code: {code}")
                    return None
            
            # Extract Quantity
            qty_match = re.search(r'QTY:\s*(\d+)', text, re.IGNORECASE)
            if qty_match:
                label_data['quantity'] = int(qty_match.group(1))
            
            # Extract Customizations
            customization_patterns = [
                r'Customizations:(.*?)(?=QTY:|$)',
                r'Left Chest Customization:(.*?)(?=QTY:|$)',
                r'Surface 1:(.*?)(?=QTY:|$)'
            ]
            
            customizations = []
            for pattern in customization_patterns:
                matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
                for match in matches:
                    customizations.append(match.strip())
            
            if customizations:
                label_data['customizations'] = '\n'.join(customizations)
            
            # Validate required fields
            if not label_data['order_id'] or not label_data.get('product'):
                print(f"Validation failed: missing Order ID or Product")
                return None
                
            return label_data
            
        except Exception as e:
            print(f"Error extracting label data: {str(e)}")
            return None


class PackingSlipsViewSet(ModelViewSet):
    """CRUD operations for packing slips"""
    queryset = PackingSlip.objects.all()
    serializer_class = PackingSlipSerializer
    permission_classes = (isAuthenticatedCustom,)

    def get_queryset(self):
        queryset = PackingSlip.objects.all()
        
        # Filter by order_id if provided
        order_id = self.request.query_params.get('order_id', None)
        if order_id:
            queryset = queryset.filter(order_id__icontains=order_id)
        
        # Filter by product code if provided
        product_code = self.request.query_params.get('product_code', None)
        if product_code:
            queryset = queryset.filter(product__code__icontains=product_code)
            
        return queryset
    
    def perform_update(self, serializer):
        """Override to call Track123 API when tracking IDs are updated"""
        instance = serializer.instance
        old_tracking_ids = instance.tracking_ids if instance else ''
        old_tracking_vendor = instance.tracking_vendor if instance else ''
        
        # Save the packing slip first
        packing_slip = serializer.save()
        
        # Check if tracking_ids or tracking_vendor were updated
        new_tracking_ids = packing_slip.tracking_ids
        new_tracking_vendor = packing_slip.tracking_vendor
        
        # If tracking info was updated and we have both tracking IDs and vendor, call Track123
        if (new_tracking_ids and new_tracking_vendor and 
            (new_tracking_ids != old_tracking_ids or new_tracking_vendor != old_tracking_vendor)):
            
            # Get Track123 API key from GoogleDriveSettings
            # Use the first active setting with an API key
            gdrive_settings = GoogleDriveSettings.objects.filter(
                is_active=True
            ).exclude(
                track123_api_key=''
            ).first()
            
            if gdrive_settings and gdrive_settings.track123_api_key:
                # Parse tracking IDs (can be comma-separated)
                tracking_numbers = [
                    tid.strip() 
                    for tid in new_tracking_ids.split(',') 
                    if tid.strip()
                ]
                
                if tracking_numbers:
                    # Call Track123 API
                    result = import_tracking_to_track123(
                        api_key=gdrive_settings.track123_api_key,
                        tracking_numbers=tracking_numbers,
                        courier_code=new_tracking_vendor
                    )
                    
                    if result.get('success'):
                        # Log successful import
                        add_user_activity(
                            self.request.user,
                            f"imported {result.get('tracking_count', len(tracking_numbers))} tracking number(s) to Track123 for order {packing_slip.order_id}"
                        )
                    else:
                        # Log error but don't fail the update
                        error_msg = result.get('error', 'Unknown error')
                        add_user_activity(
                            self.request.user,
                            f"failed to import tracking numbers to Track123 for order {packing_slip.order_id}: {error_msg}"
                        )
            else:
                # No API key configured, silently skip
                pass
    
    @action(detail=True, methods=['post'], url_path='files')
    def add_file(self, request, pk=None):
        """Manually add a file (shipping label, packing slip doc, DST, DGT) to a packing slip"""
        try:
            packing_slip = self.get_object()
            
            file_type = request.data.get('file_type')
            file_path = request.data.get('file_path')
            page_number = request.data.get('page_number')
            
            if not file_type or not file_path:
                return Response({
                    'success': False,
                    'error': 'file_type and file_path are required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Validate file type
            valid_types = ['shipping_label', 'packing_slip', 'dst', 'dgt']
            if file_type not in valid_types:
                return Response({
                    'success': False,
                    'error': f'Invalid file_type. Must be one of: {", ".join(valid_types)}'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Create file record
            file_data = {
                'packing_slip': packing_slip.id,
                'file_type': file_type,
                'file_path': file_path,
            }
            
            if page_number:
                try:
                    file_data['page_number'] = int(page_number)
                except (ValueError, TypeError):
                    pass
            
            serializer = FileSerializer(data=file_data)
            if serializer.is_valid():
                file_record = serializer.save()
                
                # Log activity
                add_user_activity(
                    request.user,
                    f"manually added {file_type} file to order {packing_slip.order_id}"
                )
                
                return Response({
                    'success': True,
                    'message': 'File added successfully',
                    'file': FileSerializer(file_record).data
                }, status=status.HTTP_201_CREATED)
            else:
                return Response({
                    'success': False,
                    'error': serializer.errors
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['delete'], url_path='files/(?P<file_id>[0-9]+)')
    def delete_file(self, request, pk=None, file_id=None):
        """Delete a file associated with this packing slip"""
        try:
            packing_slip = self.get_object()
            
            # Get the file and verify it belongs to this packing slip
            try:
                file_record = File.objects.get(id=file_id, packing_slip=packing_slip)
            except File.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'File not found or does not belong to this packing slip'
                }, status=status.HTTP_404_NOT_FOUND)
            
            file_type = file_record.get_file_type_display()
            file_record.delete()
            
            # Log activity
            add_user_activity(
                request.user,
                f"deleted {file_type} file from order {packing_slip.order_id}"
            )
            
            return Response({
                'success': True,
                'message': 'File deleted successfully'
            })
            
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PackingSlipFetchTrackingStatusView(APIView):
    """Fetch latest tracking status from Track123 and update packing slip"""
    permission_classes = (isAuthenticatedCustom,)
    
    def post(self, request, packing_slip_id):
        try:
            # Get the packing slip
            try:
                packing_slip = PackingSlip.objects.get(id=packing_slip_id)
            except PackingSlip.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'Packing slip not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Check if tracking info is available
            if not packing_slip.tracking_ids:
                return Response({
                    'success': False,
                    'error': 'No tracking IDs found for this packing slip'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not packing_slip.tracking_vendor:
                return Response({
                    'success': False,
                    'error': 'No tracking vendor specified for this packing slip'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get Track123 API key from GoogleDriveSettings
            gdrive_settings = GoogleDriveSettings.objects.filter(
                is_active=True
            ).exclude(
                track123_api_key=''
            ).first()
            
            if not gdrive_settings or not gdrive_settings.track123_api_key:
                return Response({
                    'success': False,
                    'error': 'Track123 API key is not configured. Please add it in Google Drive Settings.'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get the first tracking number (in case there are multiple)
            tracking_numbers = [
                tid.strip() 
                for tid in packing_slip.tracking_ids.split(',') 
                if tid.strip()
            ]
            
            if not tracking_numbers:
                return Response({
                    'success': False,
                    'error': 'No valid tracking numbers found'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Use the first tracking number
            tracking_number = tracking_numbers[0]
            
            # Call Track123 API to get status
            result = get_tracking_status(
                api_key=gdrive_settings.track123_api_key,
                tracking_number=tracking_number,
                courier_code=packing_slip.tracking_vendor
            )
            
            if result.get('success'):
                # Update the packing slip with the new status
                new_status = result.get('status', 'Status retrieved')
                
                # Log the full response data for debugging if status seems generic
                if 'Status retrieved' in new_status or 'see details' in new_status.lower():
                    logger.warning(f"Generic status returned. Full API response: {result.get('data')}")
                
                packing_slip.tracking_status = new_status
                packing_slip.save()
                
                # Log the activity
                add_user_activity(
                    request.user,
                    f"fetched and updated tracking status for order {packing_slip.order_id}: {new_status}"
                )
                
                return Response({
                    'success': True,
                    'message': 'Tracking status updated successfully',
                    'tracking_status': new_status,
                    'data': result.get('data')
                })
            else:
                error_msg = result.get('error', 'Unknown error')
                # Log the error
                add_user_activity(
                    request.user,
                    f"failed to fetch tracking status for order {packing_slip.order_id}: {error_msg}"
                )
                
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_400_BAD_REQUEST)
        
        except Exception as e:
            logger.error(f"Error fetching tracking status: {str(e)}")
            return Response({
                'success': False,
                'error': f'An error occurred while fetching tracking status: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ShippingLabelsUploadView(APIView):
    """Upload and process shipping labels PDF"""
    permission_classes = (isAuthenticatedCustom,)
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        try:
            file = request.FILES.get('file')
            folder_path = request.data.get('folder_path', 'Manual Upload')
            
            if not file:
                return Response({
                    'error': 'No file uploaded'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Check if file is a PDF
            if not file.name.lower().endswith('.pdf'):
                return Response({
                    'error': 'Only PDF files are supported for shipping labels'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Save the uploaded PDF temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                for chunk in file.chunks():
                    temp_file.write(chunk)
                temp_pdf_path = temp_file.name

            try:
                # Get all packing slips for matching
                packing_slips = list(PackingSlip.objects.values('id', 'ship_to', 'order_id'))
                
                # Process the PDF
                from .pdf_utils import PDFProcessor
                processor = PDFProcessor()
                processed_labels = processor.process_shipping_labels_pdf(temp_pdf_path, packing_slips)
                
                # Save shipping label files to database
                created_files = []
                unmatched_labels = []
                
                for label_data in processed_labels:
                    if label_data['matched']:
                        # Create File record for matched shipping label
                        file_record = File.objects.create(
                            packing_slip_id=label_data['packing_slip_id'],
                            file_type='shipping_label',
                            file_path=label_data['file_path']
                        )
                        created_files.append({
                            'id': file_record.id,
                            'page_number': label_data['page_number'],
                            'file_path': label_data['file_path'],
                            'packing_slip_id': label_data['packing_slip_id'],
                            'confidence_score': label_data['confidence_score']
                        })
                    else:
                        unmatched_labels.append({
                            'page_number': label_data['page_number'],
                            'shipping_address': label_data['shipping_address'],
                            'file_path': label_data['file_path']
                        })

                return Response({
                    'success': True,
                    'message': f'Processed {len(processed_labels)} shipping label(s)',
                    'matched_labels': len(created_files),
                    'unmatched_labels': len(unmatched_labels),
                    'created_files': created_files,
                    'unmatched_labels': unmatched_labels
                })

            finally:
                # Clean up temporary PDF file
                if os.path.exists(temp_pdf_path):
                    try:
                        os.unlink(temp_pdf_path)
                    except OSError:
                        pass

        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)


class FileViewerView(APIView):
    """Serve files from Google Drive through the backend"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request, file_id):
        """Get file content from Google Drive and serve it"""
        try:
            # Get the Google Drive email from request parameters or use default
            email = request.GET.get('google_drive_email')
            if not email:
                # Try to get from GoogleDriveSettings - use the first available
                from users.models import GoogleDriveSettings
                settings = GoogleDriveSettings.objects.first()
                if not settings:
                    return Response({
                        'error': 'No Google Drive settings found'
                    }, status=status.HTTP_400_BAD_REQUEST)
                email = settings.email
            
            # Initialize Google Drive service
            service = GoogleDriveService(email)
            
            # Get file info first
            file_info = service.get_file_info(file_id)
            
            # Get file content
            file_content = service.get_file_content(file_id)
            
            # Create HTTP response with file content
            from django.http import HttpResponse
            response = HttpResponse(
                file_content,
                content_type=file_info.get('mimeType', 'application/octet-stream')
            )
            
            # Set filename in response headers
            filename = file_info.get('name', f'file_{file_id}')
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            response['Content-Length'] = len(file_content)
            
            return response
            
        except Exception as e:
            return Response({
                'error': f'Error retrieving file: {str(e)}'
            }, status=status.HTTP_400_BAD_REQUEST)


class DashboardKPIView(APIView):
    """Get dashboard KPIs and analytics data"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request):
        """Get comprehensive dashboard data"""
        try:
            from django.db.models import Sum, Count, Avg, Q, F
            from django.db.models.functions import TruncMonth, TruncYear
            import datetime
            
            print("🔍 Starting dashboard KPI fetch...")
            
            # Product-level KPIs
            print("📊 Fetching product KPIs...")
            product_kpis = self.get_product_kpis()
            
            # Profit analysis by product
            print("💰 Fetching profit analysis...")
            profit_analysis = self.get_profit_analysis()
            
            # Time-based analysis
            print("📅 Fetching time analysis...")
            monthly_analysis = self.get_monthly_analysis()
            yearly_analysis = self.get_yearly_analysis()
            
            # Order status distribution
            print("📈 Fetching status distribution...")
            status_distribution = self.get_status_distribution()
            
            # Overall summary statistics
            print("📋 Fetching summary stats...")
            summary_stats = self.get_summary_statistics()
            
            print("✅ All dashboard data fetched successfully")
            
            return Response({
                'success': True,
                'data': {
                    'product_kpis': product_kpis,
                    'profit_analysis': profit_analysis,
                    'monthly_analysis': monthly_analysis,
                    'yearly_analysis': yearly_analysis,
                    'status_distribution': status_distribution,
                    'summary_stats': summary_stats
                }
            })
            
        except Exception as e:
            return Response({
                'error': f'Error retrieving dashboard data: {str(e)}'
            }, status=status.HTTP_400_BAD_REQUEST)
    
    def get_product_kpis(self):
        """Get KPIs aggregated by product"""
        from django.db.models import Sum, Count, Avg, F, Case, When, DecimalField
        
        product_data = PackingSlip.objects.values(
            'product__name', 
            'product__code',
            'product__sku_buy_cost',
            'product__sku_price'
        ).annotate(
            total_orders=Count('id'),
            total_quantity=Sum('quantity'),
            avg_sales_price=Avg('sales_price'),
            avg_shipping_price=Avg('shipping_price'),
            avg_item_cost=Avg('item_cost'),
            avg_shipping_cost=Avg('shipping_cost'),
            total_revenue=Sum(F('sales_price') + F('shipping_price')),
            total_costs=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
            total_profit=Sum('profit'),
            avg_profit_margin=Avg(
                F('profit') / (F('sales_price') + F('shipping_price')) * 100
            )
        ).order_by('-total_profit')
        
        return list(product_data)
    
    def get_profit_analysis(self):
        """Get detailed profit analysis"""
        from django.db.models import Sum, Count, Avg, F, Case, When, DecimalField
        
        # Top performing products by profit margin
        top_profit_margin = PackingSlip.objects.values(
            'product__name', 
            'product__code'
        ).annotate(
            avg_profit_margin=Avg(
                F('profit') / (F('sales_price') + F('shipping_price')) * 100
            ),
            total_orders=Count('id'),
            total_profit=Sum('profit')
        ).filter(
            avg_profit_margin__isnull=False
        ).order_by('-avg_profit_margin')[:10]
        
        # Profit distribution by ranges
        profit_ranges = {
            'high_profit': PackingSlip.objects.filter(profit__gte=50).count(),
            'medium_profit': PackingSlip.objects.filter(profit__gte=20, profit__lt=50).count(),
            'low_profit': PackingSlip.objects.filter(profit__gte=0, profit__lt=20).count(),
            'loss_making': PackingSlip.objects.filter(profit__lt=0).count()
        }
        
        return {
            'top_profit_margin_products': list(top_profit_margin),
            'profit_distribution': profit_ranges
        }
    
    def get_monthly_analysis(self):
        """Get monthly trend analysis"""
        from django.db.models import Sum, Count, Avg, F
        from django.db.models.functions import TruncMonth
        
        monthly_data = PackingSlip.objects.annotate(
            month=TruncMonth('created_at')
        ).values('month').annotate(
            total_orders=Count('id'),
            total_revenue=Sum(F('sales_price') + F('shipping_price')),
            total_profit=Sum('profit'),
            avg_profit_margin=Avg(
                F('profit') / (F('sales_price') + F('shipping_price')) * 100
            )
        ).order_by('month')
        
        return list(monthly_data)
    
    def get_yearly_analysis(self):
        """Get yearly trend analysis"""
        from django.db.models import Sum, Count, Avg, F
        from django.db.models.functions import TruncYear
        
        yearly_data = PackingSlip.objects.annotate(
            year=TruncYear('created_at')
        ).values('year').annotate(
            total_orders=Count('id'),
            total_revenue=Sum(F('sales_price') + F('shipping_price')),
            total_profit=Sum('profit'),
            avg_profit_margin=Avg(
                F('profit') / (F('sales_price') + F('shipping_price')) * 100
            )
        ).order_by('year')
        
        return list(yearly_data)
    
    def get_status_distribution(self):
        """Get order distribution by status"""
        from django.db.models import Count, Sum, F
        
        status_data = PackingSlip.objects.values('status').annotate(
            count=Count('id'),
            total_value=Sum(F('sales_price') + F('shipping_price'))
        ).order_by('-count')
        
        # Add readable status names
        status_mapping = dict(PackingSlip.STATUS_CHOICES)
        for item in status_data:
            item['status_display'] = status_mapping.get(item['status'], item['status'])
        
        return list(status_data)
    
    def get_summary_statistics(self):
        """Get overall summary statistics"""
        from django.db.models import Sum, Count, Avg, F
        from .models import Expense
        
        total_orders = PackingSlip.objects.count()
        total_revenue = PackingSlip.objects.aggregate(
            total=Sum(F('sales_price') + F('shipping_price'))
        )['total'] or 0
        
        # Get total expenses from Expense model
        total_expenses = Expense.objects.aggregate(
            total=Sum('amount')
        )['total'] or 0
        
        # Calculate profit from packing slips
        packing_slip_profit = PackingSlip.objects.aggregate(
            total=Sum('profit')
        )['total'] or 0
        
        # Total profit = packing slip profit - total expenses
        total_profit = float(packing_slip_profit) - float(total_expenses)
        
        avg_order_value = PackingSlip.objects.aggregate(
            avg=Avg(F('sales_price') + F('shipping_price'))
        )['avg'] or 0
        
        return {
            'total_orders': total_orders,
            'total_revenue': float(total_revenue),
            'total_profit': total_profit,
            'total_expenses': float(total_expenses),
            'avg_order_value': float(avg_order_value)
        }


class DashboardCustomerProfitAnalysisView(APIView):
    """Get customer profit analysis from folder paths"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request):
        """Get customer profit analysis data"""
        try:
            from django.db.models import Sum, Count, Avg, Q, F
            import re
            
            print("🔍 Starting customer profit analysis...")
            
            # Get all packing slips with folder paths
            packing_slips = PackingSlip.objects.exclude(
                Q(folder_path__isnull=True) | Q(folder_path__exact='') | Q(folder_path__exact='Manual Upload')
            ).values('folder_path', 'profit', 'sales_price', 'shipping_price')
            
            print(f"Found {len(packing_slips)} packing slips with folder paths")
            
            # Dictionary to store customer profit data
            customer_data = {}
            
            # Process each packing slip to extract customer from folder path
            for slip in packing_slips:
                folder_path = slip['folder_path']
                customer = self.extract_customer_from_path(folder_path)
                
                if customer and customer != 'Unknown':
                    if customer not in customer_data:
                        customer_data[customer] = {
                            'customer': customer,
                            'total_profit': 0,
                            'total_revenue': 0,
                            'order_count': 0
                        }
                    
                    # Add profit and revenue data
                    profit = float(slip['profit'] or 0)
                    revenue = float((slip['sales_price'] or 0) + (slip['shipping_price'] or 0))
                    
                    customer_data[customer]['total_profit'] += profit
                    customer_data[customer]['total_revenue'] += revenue
                    customer_data[customer]['order_count'] += 1
            
            # Get total expenses to deduct from overall profit
            from .models import Expense
            total_expenses = Expense.objects.aggregate(
                total=Sum('amount')
            )['total'] or 0
            
            # Calculate total profit and revenue
            total_customer_profit = sum(c['total_profit'] for c in customer_data.values())
            total_customer_revenue = sum(c['total_revenue'] for c in customer_data.values())
            
            # Deduct expenses proportionally from each customer based on their revenue share
            for customer in customer_data.values():
                if total_customer_revenue > 0:
                    # Calculate this customer's share of total expenses
                    revenue_share = customer['total_revenue'] / total_customer_revenue
                    customer_expense_share = float(total_expenses) * revenue_share
                    
                    # Deduct expenses from profit
                    customer['total_profit'] -= customer_expense_share
                    customer['expense_allocation'] = customer_expense_share
            
            # Convert to list and sort by profit
            customer_profit_data = list(customer_data.values())
            customer_profit_data.sort(key=lambda x: x['total_profit'], reverse=True)
            
            print(f"✅ Customer profit analysis completed for {len(customer_profit_data)} customers")
            print(f"💰 Total expenses allocated: ${float(total_expenses):.2f}")
            
            return Response({
                'success': True,
                'data': {
                    'customer_profit_analysis': customer_profit_data,
                    'total_expenses': float(total_expenses)
                }
            })
            
        except Exception as e:
            print(f"Error in customer profit analysis: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({
                'error': f'Error retrieving customer profit analysis: {str(e)}'
            }, status=status.HTTP_400_BAD_REQUEST)
    
    def extract_customer_from_path(self, folder_path):
        """Extract customer name from folder path like 'EMB / 09082025 / Amazon / Packing Slips'"""
        try:
            if not folder_path:
                return 'Unknown'
            
            # Split path by " / " separator
            path_parts = [part.strip() for part in folder_path.split('/')]
            
            # Expected format: EMB / DATE / CUSTOMER / SUBFOLDER
            if len(path_parts) >= 3:
                # The customer should be the third part (index 2)
                customer = path_parts[2].strip()
                
                # Clean up common variations
                if customer and customer.lower() not in ['emb', 'packing slips', 'shipping labels', 'dst', 'pick lists']:
                    return customer
            
            return 'Unknown'
            
        except Exception as e:
            print(f"Error extracting customer from path '{folder_path}': {e}")
            return 'Unknown'


class DashboardSKUAnalysisView(APIView):
    """Get SKU-level dashboard analysis with filtering"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request):
        """Get SKU-level analysis data"""
        try:
            from django.db.models import Sum, Count, Avg, Q, F
            from django.db.models.functions import TruncMonth, TruncYear
            import datetime
            
            print("🔍 Starting SKU-level dashboard analysis...")
            
            # Get query parameters for filtering
            selected_skus = request.GET.getlist('skus[]', [])  # List of selected SKU codes
            time_period = request.GET.get('time_period', 'all')  # all, last_6_months, last_3_months, last_year
            
            print(f"Selected SKUs: {selected_skus}")
            print(f"Time period: {time_period}")
            
            # Base queryset
            base_queryset = PackingSlip.objects.all()
            
            # Apply time filtering
            if time_period == 'last_3_months':
                cutoff_date = datetime.datetime.now() - datetime.timedelta(days=90)
                base_queryset = base_queryset.filter(created_at__gte=cutoff_date)
            elif time_period == 'last_6_months':
                cutoff_date = datetime.datetime.now() - datetime.timedelta(days=180)
                base_queryset = base_queryset.filter(created_at__gte=cutoff_date)
            elif time_period == 'last_year':
                cutoff_date = datetime.datetime.now() - datetime.timedelta(days=365)
                base_queryset = base_queryset.filter(created_at__gte=cutoff_date)
            
            # Apply SKU filtering
            if selected_skus:
                base_queryset = base_queryset.filter(product__code__in=selected_skus)
            
            # Get available SKUs for filtering
            available_skus = list(Product.objects.values('code', 'name').order_by('name'))
            
            # Monthly analysis with SKU breakdown
            print("📅 Fetching monthly SKU analysis...")
            monthly_analysis = self.get_monthly_sku_analysis(base_queryset, selected_skus)
            
            # Yearly analysis with SKU breakdown
            print("📊 Fetching yearly SKU analysis...")
            yearly_analysis = self.get_yearly_sku_analysis(base_queryset, selected_skus)
            
            # SKU performance summary
            print("💰 Fetching SKU performance summary...")
            sku_summary = self.get_sku_performance_summary(base_queryset)
            
            print("✅ SKU-level analysis completed successfully")
            
            return Response({
                'success': True,
                'data': {
                    'available_skus': available_skus,
                    'selected_skus': selected_skus,
                    'time_period': time_period,
                    'monthly_analysis': monthly_analysis,
                    'yearly_analysis': yearly_analysis,
                    'sku_summary': sku_summary
                }
            })
            
        except Exception as e:
            print(f"Error in SKU analysis: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({
                'error': f'Error retrieving SKU analysis: {str(e)}'
            }, status=status.HTTP_400_BAD_REQUEST)
    
    def get_monthly_sku_analysis(self, base_queryset, selected_skus):
        """Get monthly analysis with SKU-level breakdown"""
        from django.db.models import Sum, Count, Avg, F
        from django.db.models.functions import TruncMonth
        
        if selected_skus:
            # Individual SKU breakdown by month
            monthly_data = base_queryset.annotate(
                month=TruncMonth('created_at')
            ).values('month', 'product__code', 'product__name').annotate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            ).order_by('month', 'product__code')
            
            # Group by month and aggregate across selected SKUs
            monthly_aggregated = {}
            for item in monthly_data:
                month_key = item['month'].strftime('%Y-%m-%d')
                if month_key not in monthly_aggregated:
                    monthly_aggregated[month_key] = {
                        'month': item['month'],
                        'total_orders': 0,
                        'total_revenue': 0,
                        'total_profit': 0,
                        'total_cost': 0,
                        'skus': []
                    }
                
                monthly_aggregated[month_key]['total_orders'] += item['total_orders'] or 0
                monthly_aggregated[month_key]['total_revenue'] += item['total_revenue'] or 0
                monthly_aggregated[month_key]['total_profit'] += item['total_profit'] or 0
                monthly_aggregated[month_key]['total_cost'] += item['total_cost'] or 0
                monthly_aggregated[month_key]['skus'].append({
                    'sku_code': item['product__code'],
                    'sku_name': item['product__name'],
                    'orders': item['total_orders'],
                    'revenue': float(item['total_revenue'] or 0),
                    'profit': float(item['total_profit'] or 0),
                    'cost': float(item['total_cost'] or 0)
                })
            
            # Convert to list and add profit margin
            result = []
            for month_data in monthly_aggregated.values():
                if month_data['total_revenue'] > 0:
                    month_data['avg_profit_margin'] = (month_data['total_profit'] / month_data['total_revenue']) * 100
                else:
                    month_data['avg_profit_margin'] = 0
                result.append(month_data)
            
            return sorted(result, key=lambda x: x['month'])
            
        else:
            # Aggregate all SKUs by month
            monthly_data = base_queryset.annotate(
                month=TruncMonth('created_at')
            ).values('month').annotate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            ).order_by('month')
            
            return [
                {
                    'month': item['month'],
                    'total_orders': item['total_orders'],
                    'total_revenue': float(item['total_revenue'] or 0),
                    'total_profit': float(item['total_profit'] or 0),
                    'total_cost': float(item['total_cost'] or 0),
                    'avg_profit_margin': float(item['avg_profit_margin'] or 0),
                    'skus': []  # Empty for aggregated view
                }
                for item in monthly_data
            ]
    
    def get_yearly_sku_analysis(self, base_queryset, selected_skus):
        """Get yearly analysis with SKU-level breakdown"""
        from django.db.models import Sum, Count, Avg, F
        from django.db.models.functions import TruncYear
        
        if selected_skus:
            # Individual SKU breakdown by year
            yearly_data = base_queryset.annotate(
                year=TruncYear('created_at')
            ).values('year', 'product__code', 'product__name').annotate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            ).order_by('year', 'product__code')
            
            # Group by year and aggregate across selected SKUs
            yearly_aggregated = {}
            for item in yearly_data:
                year_key = item['year'].year
                if year_key not in yearly_aggregated:
                    yearly_aggregated[year_key] = {
                        'year': year_key,
                        'total_orders': 0,
                        'total_revenue': 0,
                        'total_profit': 0,
                        'total_cost': 0,
                        'skus': []
                    }
                
                yearly_aggregated[year_key]['total_orders'] += item['total_orders'] or 0
                yearly_aggregated[year_key]['total_revenue'] += item['total_revenue'] or 0
                yearly_aggregated[year_key]['total_profit'] += item['total_profit'] or 0
                yearly_aggregated[year_key]['total_cost'] += item['total_cost'] or 0
                yearly_aggregated[year_key]['skus'].append({
                    'sku_code': item['product__code'],
                    'sku_name': item['product__name'],
                    'orders': item['total_orders'],
                    'revenue': float(item['total_revenue'] or 0),
                    'profit': float(item['total_profit'] or 0),
                    'cost': float(item['total_cost'] or 0)
                })
            
            # Convert to list and add profit margin
            result = []
            for year_data in yearly_aggregated.values():
                if year_data['total_revenue'] > 0:
                    year_data['avg_profit_margin'] = (year_data['total_profit'] / year_data['total_revenue']) * 100
                else:
                    year_data['avg_profit_margin'] = 0
                result.append(year_data)
            
            return sorted(result, key=lambda x: x['year'])
            
        else:
            # Aggregate all SKUs by year
            yearly_data = base_queryset.annotate(
                year=TruncYear('created_at')
            ).values('year').annotate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            ).order_by('year')
            
            return [
                {
                    'year': item['year'].year,
                    'total_orders': item['total_orders'],
                    'total_revenue': float(item['total_revenue'] or 0),
                    'total_profit': float(item['total_profit'] or 0),
                    'total_cost': float(item['total_cost'] or 0),
                    'avg_profit_margin': float(item['avg_profit_margin'] or 0),
                    'skus': []  # Empty for aggregated view
                }
                for item in yearly_data
            ]
    
    def get_sku_performance_summary(self, base_queryset):
        """Get performance summary for each SKU"""
        from django.db.models import Sum, Count, Avg, F
        
        sku_data = base_queryset.values(
            'product__code', 
            'product__name'
        ).annotate(
            total_orders=Count('id'),
            total_revenue=Sum(F('sales_price') + F('shipping_price')),
            total_profit=Sum('profit'),
            total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
            avg_profit_margin=Avg(
                F('profit') / (F('sales_price') + F('shipping_price')) * 100
            )
        ).order_by('-total_profit')
        
        return [
            {
                'sku_code': item['product__code'],
                'sku_name': item['product__name'],
                'total_orders': item['total_orders'],
                'total_revenue': float(item['total_revenue'] or 0),
                'total_profit': float(item['total_profit'] or 0),
                'total_cost': float(item['total_cost'] or 0),
                'avg_profit_margin': float(item['avg_profit_margin'] or 0)
            }
            for item in sku_data
        ]


class OrdersByStatusView(APIView):
    """Get orders filtered by status"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request):
        try:
            status_filter = request.GET.get('status')
            
            if not status_filter:
                return Response({
                    'success': False,
                    'error': 'Status parameter is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get orders with the specified status
            orders = PackingSlip.objects.filter(status=status_filter).select_related('product').order_by('-created_at')
            
            # Format the data for the frontend
            orders_data = []
            for order in orders:
                # Get status display name
                status_display = dict(PackingSlip.STATUS_CHOICES).get(order.status, order.status)
                
                # Calculate total price (sales_price + shipping_price)
                total_price = float((order.sales_price or 0) + (order.shipping_price or 0))
                
                orders_data.append({
                    'id': order.id,
                    'order_id': order.order_id,
                    'asin': order.asin,
                    'product__name': order.product.name,
                    'product__code': order.product.code,
                    'customer_name': order.ship_to.split('\n')[0] if order.ship_to else 'N/A',  # First line as customer name
                    'quantity': order.quantity,
                    'status': order.status,
                    'status_display': status_display,
                    'total_price': total_price,
                    'sales_price': float(order.sales_price or 0),
                    'shipping_price': float(order.shipping_price or 0),
                    'profit': float(order.profit or 0),
                    'customizations': order.customizations,
                    'created_at': order.created_at.isoformat(),
                    'updated_at': order.updated_at.isoformat(),
                })
            
            return Response({
                'success': True,
                'data': orders_data,
                'count': len(orders_data)
            })
            
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ProductAnalyticsView(APIView):
    """Get detailed analytics for a specific product"""
    permission_classes = (isAuthenticatedCustom,)
    
    def get(self, request):
        try:
            product_code = request.GET.get('product_code')
            
            if not product_code:
                return Response({
                    'success': False,
                    'error': 'Product code parameter is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Check if product exists
            try:
                product = Product.objects.get(code=product_code)
            except Product.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'Product not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Get all orders for this product
            orders = PackingSlip.objects.filter(product=product).order_by('-created_at')
            
            if not orders.exists():
                return Response({
                    'success': True,
                    'data': {
                        'product': {
                            'code': product.code,
                            'name': product.name,
                            'sku_price': float(product.sku_price),
                            'sku_buy_cost': float(product.sku_buy_cost)
                        },
                        'summary': {
                            'total_orders': 0,
                            'total_revenue': 0,
                            'total_profit': 0,
                            'total_cost': 0,
                            'avg_profit_margin': 0
                        },
                        'monthly_data': []
                    }
                })
            
            from django.db.models import Sum, Count, Avg, F
            from django.db.models.functions import TruncMonth
            
            # Calculate summary statistics
            summary_stats = orders.aggregate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            )
            
            # Get monthly breakdown
            monthly_data = orders.annotate(
                month=TruncMonth('created_at')
            ).values('month').annotate(
                total_orders=Count('id'),
                total_revenue=Sum(F('sales_price') + F('shipping_price')),
                total_profit=Sum('profit'),
                total_cost=Sum(F('item_cost') + F('shipping_cost') + F('platform_fee_calculated')),
                avg_profit_margin=Avg(
                    F('profit') / (F('sales_price') + F('shipping_price')) * 100
                )
            ).order_by('month')
            
            # Format monthly data
            monthly_formatted = [
                {
                    'month': item['month'].isoformat(),
                    'total_orders': item['total_orders'],
                    'total_revenue': float(item['total_revenue'] or 0),
                    'total_profit': float(item['total_profit'] or 0),
                    'total_cost': float(item['total_cost'] or 0),
                    'avg_profit_margin': float(item['avg_profit_margin'] or 0)
                }
                for item in monthly_data
            ]
            
            return Response({
                'success': True,
                'data': {
                    'product': {
                        'code': product.code,
                        'name': product.name,
                        'sku_price': float(product.sku_price),
                        'sku_buy_cost': float(product.sku_buy_cost),
                        'description': product.sku_description,
                        'uom': product.sku_uom,
                        'color': product.color
                    },
                    'summary': {
                        'total_orders': summary_stats['total_orders'],
                        'total_revenue': float(summary_stats['total_revenue'] or 0),
                        'total_profit': float(summary_stats['total_profit'] or 0),
                        'total_cost': float(summary_stats['total_cost'] or 0),
                        'avg_profit_margin': float(summary_stats['avg_profit_margin'] or 0)
                    },
                    'monthly_data': monthly_formatted
                }
            })
            
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PackingSlipBulkPrintView(APIView):
    """Generate bulk PDF with multiple packing slips"""
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        try:
            # Get the list of packing slip IDs
            packing_slip_ids = request.data.get('packing_slip_ids', [])
            
            if not packing_slip_ids:
                return Response({
                    'error': 'No packing slip IDs provided'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get all packing slips
            packing_slips = PackingSlip.objects.filter(id__in=packing_slip_ids)
            
            if not packing_slips.exists():
                return Response({
                    'error': 'No packing slips found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.5*inch, 
                                  leftMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            # Container for all 'Flowable' objects
            elements = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Process each packing slip
            for i, packing_slip in enumerate(packing_slips):
                print(f"Processing packing slip {i+1}/{len(packing_slips)}: Order {packing_slip.order_id}")
                
                # Add shipping label page first (if available)
                shipping_labels = packing_slip.files.filter(file_type='shipping_label')
                if shipping_labels.exists():
                    # Use the existing single print view method
                    single_print_view = PackingSlipPrintView()
                    elements.extend(single_print_view.add_shipping_label_page(shipping_labels.first(), styles))
                    # Add page break after shipping label
                    from reportlab.platypus import PageBreak
                    elements.append(PageBreak())
                
                # Add packing slip page
                single_print_view = PackingSlipPrintView()
                elements.extend(single_print_view.create_packing_slip_page(packing_slip, styles))
                
                # Add page break between different orders (except for the last one)
                if i < len(packing_slips) - 1:
                    from reportlab.platypus import PageBreak
                    elements.append(PageBreak())
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            # Create response
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="bulk_packing_slips_{len(packing_slips)}_orders.pdf"'
            
            return response
            
        except Exception as e:
            print(f"Error in bulk print: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PackingSlipPrintView(APIView):
    """Generate PDF with shipping label and packing slip"""
    permission_classes = (isAuthenticatedCustom,)

    def get(self, request, packing_slip_id):
        try:
            # Get the packing slip
            packing_slip = PackingSlip.objects.get(id=packing_slip_id)
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.5*inch, 
                                  leftMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Add shipping label page first (if available)
            shipping_labels = packing_slip.files.filter(file_type='shipping_label')
            if shipping_labels.exists():
                elements.extend(self.add_shipping_label_page(shipping_labels.first(), styles))
                # Add page break after shipping label
                from reportlab.platypus import PageBreak
                elements.append(PageBreak())
            
            # Add packing slip page
            elements.extend(self.create_packing_slip_page(packing_slip, styles))
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            # Create response
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="packing_slip_{packing_slip.order_id}.pdf"'
            
            return response
            
        except PackingSlip.DoesNotExist:
            return Response({
                'error': 'Packing slip not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def add_shipping_label_page(self, shipping_label_file, styles):
        """Add shipping label page from PDF using Google Drive service"""
        elements = []
        
        try:
            print(f"Processing shipping label file: {shipping_label_file.file_path}")
            print(f"Page number: {shipping_label_file.page_number}")
            
            # Extract file ID from Google Drive URL
            file_url = shipping_label_file.file_path
            file_id = None
            
            if 'drive.google.com' in file_url:
                if '/file/d/' in file_url:
                    file_id = file_url.split('/file/d/')[1].split('/')[0]
                elif 'id=' in file_url:
                    file_id = file_url.split('id=')[1].split('&')[0]
            
            if not file_id:
                raise Exception(f"Could not extract file ID from URL: {file_url}")
            
            print(f"Extracted file ID: {file_id}")
            
            # We need to get a Google Drive service instance
            # For now, let's try to get the first available drive account
            from users.models import GoogleDriveSettings
            drive_settings = GoogleDriveSettings.objects.filter(is_active=True).first()
            
            if not drive_settings:
                raise Exception("No active Google Drive settings found")
            
            print(f"Using Google Drive account: {drive_settings.email}")
            
            # Initialize Google Drive service
            drive_service = GoogleDriveService(drive_settings.email)
            
            # Download PDF content
            pdf_content = drive_service.get_file_content(file_id)
            print(f"Downloaded PDF content: {len(pdf_content)} bytes")
            
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
            print(f"PDF has {pdf_document.page_count} pages")
            
            # Get the specified page (page_number is 1-based, PyMuPDF is 0-based)
            page_num = (shipping_label_file.page_number or 1) - 1
            if page_num >= pdf_document.page_count:
                page_num = 0  # Default to first page if specified page doesn't exist
                
            print(f"Extracting page {page_num + 1}")
            page = pdf_document[page_num]
            
            # Convert PDF page to image
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Create PIL image from PNG data
            pil_img = PILImage.open(BytesIO(img_data))
            print(f"Converted to image: {pil_img.size}")
            
            # Calculate dimensions to fit page while maintaining aspect ratio
            # Use more conservative margins to ensure image fits
            page_width = letter[0] - 1.5*inch  # Leave more margins
            page_height = letter[1] - 1.5*inch
            
            img_width, img_height = pil_img.size
            aspect_ratio = img_width / img_height
            
            print(f"Original image size: {img_width} x {img_height}")
            print(f"Available space: {page_width} x {page_height}")
            print(f"Aspect ratio: {aspect_ratio}")
            
            # Calculate scaling to fit within both width and height constraints
            width_scale = page_width / img_width
            height_scale = page_height / img_height
            scale = min(width_scale, height_scale)  # Use the smaller scale to ensure it fits
            
            new_width = img_width * scale
            new_height = img_height * scale
            
            print(f"Calculated new size: {new_width} x {new_height}")
            
            # Ensure dimensions don't exceed page size (safety check)
            if new_width > page_width:
                new_width = page_width
                new_height = page_width / aspect_ratio
            if new_height > page_height:
                new_height = page_height
                new_width = page_height * aspect_ratio
                
            print(f"Final size after safety check: {new_width} x {new_height}")
            
            # Convert to JPEG for ReportLab
            img_buffer = BytesIO()
            if pil_img.mode in ('RGBA', 'P'):
                pil_img = pil_img.convert('RGB')
            pil_img.save(img_buffer, format='JPEG', quality=85)
            img_buffer.seek(0)
            
            # Create ReportLab Image
            img = Image(img_buffer, width=new_width, height=new_height)
            elements.append(img)
            
            # Clean up
            pdf_document.close()
            print("Successfully added shipping label page to PDF")
                
        except Exception as e:
            print(f"Error loading shipping label: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # If image fails to load, add a placeholder with debug info
            error_style = ParagraphStyle('Error', parent=styles['Normal'], textColor=colors.red)
            elements.append(Paragraph(f"Shipping Label Not Available", error_style))
            elements.append(Paragraph(f"URL: {shipping_label_file.file_path}", error_style))
            elements.append(Paragraph(f"Page: {shipping_label_file.page_number or 'Not specified'}", error_style))
            elements.append(Paragraph(f"Error: {str(e)}", error_style))
            elements.append(Spacer(1, 0.2*inch))
        
        return elements

    def create_packing_slip_page(self, packing_slip, styles):
        """Create packing slip page with specified format"""
        elements = []
        
        # Custom styles - Increased font sizes
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], 
                                    fontSize=24, spaceAfter=20, alignment=TA_CENTER)
        
        header_style = ParagraphStyle('Header', parent=styles['Normal'], 
                                    fontSize=16, fontName='Helvetica-Bold', spaceAfter=8)
        
        content_style = ParagraphStyle('Content', parent=styles['Normal'], 
                                     fontSize=14, spaceAfter=6)
        
        # Title
        elements.append(Paragraph("PACKING SLIP", title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Ship To Section
        elements.append(Paragraph("Ship to", header_style))
        
        # Format ship_to address
        ship_to_lines = packing_slip.ship_to.split('\n') if packing_slip.ship_to else ["No shipping address provided"]
        for line in ship_to_lines:
            if line.strip():
                elements.append(Paragraph(line.strip(), content_style))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Order ID
        elements.append(Paragraph(f"Order ID: # {packing_slip.order_id}", header_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # ASIN
        elements.append(Paragraph(f"ASIN: {packing_slip.asin}", content_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # SKU
        elements.append(Paragraph(f"SKU: {packing_slip.product.code}", content_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Customizations
        if packing_slip.customizations:
            elements.append(Paragraph("Customizations:", header_style))
            
            # Parse customizations (assuming they might be in different formats)
            customization_lines = packing_slip.customizations.split('\n')
            for line in customization_lines:
                if line.strip():
                    elements.append(Paragraph(line.strip(), content_style))
        else:
            elements.append(Paragraph("Customizations: None", content_style))
        
        elements.append(Spacer(1, 0.1*inch))
        
        # Quantity
        qty_text = f"Qty: {packing_slip.quantity:02d}"
        elements.append(Paragraph(qty_text, header_style))

        return elements


class TrackingSchedulerView(APIView):
    """
    Manage the tracking status update scheduler.

    POST /api/masterdata/tracking/scheduler/
        - action: 'start' | 'stop' | 'status' | 'run_now'
        - interval: int (minutes, default: 720 / 12 hours) - only used with 'start' action
    """
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        from masterdata.tracking_service import (
            schedule_tracking_updates,
            cancel_tracking_schedule,
            trigger_immediate_update
        )
        from django_q.models import Schedule

        action = request.data.get('action')

        if not action:
            return Response({
                'success': False,
                'error': 'Action is required (start, stop, status, or run_now)'
            }, status=status.HTTP_400_BAD_REQUEST)

        if action == 'start':
            interval = request.data.get('interval', 720)
            try:
                interval = int(interval)
                if interval < 1:
                    raise ValueError("Interval must be at least 1 minute")
            except (ValueError, TypeError) as e:
                return Response({
                    'success': False,
                    'error': f'Invalid interval: {str(e)}'
                }, status=status.HTTP_400_BAD_REQUEST)

            result = schedule_tracking_updates(interval_minutes=interval)
            return Response(result)

        elif action == 'stop':
            result = cancel_tracking_schedule()
            return Response(result)

        elif action == 'status':
            schedules = Schedule.objects.filter(
                func='masterdata.tracking_service.update_all_pending_orders',
                name='tracking_status_updater'
            )

            if schedules.exists():
                schedule_info = []
                for schedule in schedules:
                    schedule_info.append({
                        'id': schedule.id,
                        'name': schedule.name,
                        'interval_minutes': schedule.minutes,
                        'next_run': schedule.next_run,
                        'repeats': 'Indefinitely' if schedule.repeats == -1 else schedule.repeats,
                        'task_count': schedule.task_count()
                    })

                return Response({
                    'success': True,
                    'active': True,
                    'schedules': schedule_info
                })
            else:
                return Response({
                    'success': True,
                    'active': False,
                    'message': 'Scheduler is not active'
                })

        elif action == 'run_now':
            result = trigger_immediate_update()
            return Response(result)

        else:
            return Response({
                'success': False,
                'error': 'Invalid action. Must be: start, stop, status, or run_now'
            }, status=status.HTTP_400_BAD_REQUEST)


class TrackingUpdateView(APIView):
    """
    Update tracking status for specific order(s).

    POST /api/masterdata/tracking/update/
        - order_id: string (optional) - Update specific order
        - packing_slip_id: int (optional) - Update specific packing slip
        - If neither provided, updates all pending orders
    """
    permission_classes = (isAuthenticatedCustom,)

    def post(self, request):
        from masterdata.tracking_service import update_single_order_tracking, update_all_pending_orders
        from masterdata.models import PackingSlip

        order_id = request.data.get('order_id')
        packing_slip_id = request.data.get('packing_slip_id')

        # Update specific order by packing_slip_id
        if packing_slip_id:
            try:
                packing_slip_id = int(packing_slip_id)
                result = update_single_order_tracking(packing_slip_id)
                return Response(result)
            except (ValueError, TypeError):
                return Response({
                    'success': False,
                    'error': 'Invalid packing_slip_id'
                }, status=status.HTTP_400_BAD_REQUEST)

        # Update specific order by order_id
        elif order_id:
            try:
                packing_slip = PackingSlip.objects.get(order_id=order_id)
                result = update_single_order_tracking(packing_slip.id)
                return Response(result)
            except PackingSlip.DoesNotExist:
                return Response({
                    'success': False,
                    'error': f'Order {order_id} not found'
                }, status=status.HTTP_404_NOT_FOUND)

        # Update all pending orders
        else:
            result = update_all_pending_orders()
            return Response(result)
